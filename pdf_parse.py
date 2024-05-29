import os
import sys
import time

# import requests
import asyncio
from asyncio import run

import aiohttp
from bs4 import BeautifulSoup
from dataclasses import dataclass, field, asdict
from typing import List, Optional, Any
import json
import csv
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from langdetect import detect_langs, LangDetectException
from docx import Document
import regex as re

load_dotenv()

CSV_COURSES = os.getenv("CSV_COURSES")
CSV_DEPARTMENTS = os.getenv("CSV_DEPARTMENTS")
FACULTY_NAME = os.getenv("FACULTY_NAME")
FACULTY_LINK = os.getenv("FACULTY_LINK")

@dataclass
class Teacher:
    name: str
    link: str


@dataclass
class Specialty:
    title: str
    link: str
    lecturers: Optional[List[Teacher]] = None


@dataclass
class LangStatistics:
    ukrainian_lit_count: int = 0
    moscowian_lit_count: int = 0
    other_lit_count: int = 0


@dataclass
class Course:
    title: str
    link: str
    lecturers: List[Teacher] = field(default_factory=list)
    specialties: Optional[List[Specialty]] = None
    syllabus_links: List[str] = field(default_factory=list)
    literature_statistics: Optional[LangStatistics] = None


@dataclass
class Department:
    name: str
    head: str
    phone: str
    email: str
    link: str
    course_list: List[Course] = field(default_factory=list)


@dataclass
class Faculty:
    name: str
    departments: List[Department] = field(default_factory=list)


# Function to detect language
def detect_language(text):
    try:
        probabilities = detect_langs(text)
        return probabilities[0].lang if probabilities else None
    except LangDetectException as e:
        print(f"Language detection failed: {e}")
        return None  # or return a default like 'unknown'


# Async HTTP request with aiohttp
async def fetch_html(url: str):
    async with aiohttp.ClientSession() as session:
        async with session.get(url) as response:
            return await response.text()


# Main async function to process departments
async def process_department_response(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    departments = []
    sections = soup.find_all('section')
    print(f"Found {len(sections)} sections")

    for section in sections:
        try:
            department = await parse_department(section)
            if department:
                departments.append(department)
                print(f"Added department: {department.name}")
        except Exception as e:
            print(f"Failed to parse department: {str(e)}")

    return Faculty(name="Faculty of Applied Mathematics", departments=departments)


# Function to parse department from section
async def parse_department(section):
    title_element = section.select_one('h2.title a')
    if not title_element:
        print("No title element found in section")  # Debug: No title element
        return None

    if title_element:
        name = title_element.text.strip() if title_element else "Unknown Department"
        head = section.select_one('p span.icon.user + a')
        phone = section.select_one('p span.icon.phone + a')
        email = section.select_one('p span.icon.email + a')

        # Extract text only if elements are found, otherwise use a placeholder
        head_text = head.text.strip() if head else "No Head Available"
        phone_text = phone.text.strip() if phone else "No Phone Available"
        email_text = email.text.strip() if email else "No Email Available"
        link = title_element['href'] if title_element else ""

        print(f"Parsing department: {name}")  # Debug: Check which department is parsed

        courses = await fetch_courses(link)
        return Department(name=name, head=head, phone=phone, email=email, link=link, course_list=courses)
    else:
        print("No title element found in section")  # Debug: No title element
        return None


# Async function to fetch courses and their details
async def fetch_courses(link):
    html_content = await fetch_html(link)
    soup = BeautifulSoup(html_content, 'html.parser')
    courses = []

    for a_tag in soup.select('div.item a'):
        course_url = a_tag['href']
        lecturers, syllabus_links = await fetch_course_details(course_url)
        literature_list = await fetch_literature_list(syllabus_links)
        # Aggregate literature statistics based on detected languages
        lang_stats = LangStatistics()
        for entry in literature_list:
            lang = detect_language(entry)
            if lang == 'uk':
                lang_stats.ukrainian_lit_count += 1
            elif lang == 'ru':
                lang_stats.moscowian_lit_count += 1
            elif lang is not None:  # Only increment 'other' if the language was detected but is not 'uk' or 'ru'
                lang_stats.other_lit_count += 1

        course = Course(
            title=a_tag.text.strip(),
            link=course_url,
            lecturers=lecturers,
            syllabus_links=syllabus_links,
            literature_statistics=lang_stats
        )
        courses.append(course)
    return courses


# Process extracted text to find literature entries
def process_extracted_text(text):
    patterns = [
        r"\d+\.\s+\p{L}(?:[\p{L}\s\d\'\.,\-:\—\«\»\(\)\\\/]|(?!\d\.\s+\p{L}))+",
        r"\d+\.\p{L}(?:[\p{L}\s\d\'\.,\-:\—\«\»\(\)\\\/]|(?!\d\.\p{L}))+",
        r"•\s+[\p{L}\s\d\'\.,\-:\—\«\»\(\)\\\/\"]+",
        r"•[\p{L}\s\d\'\.,\-:\—\«\»\(\)\\\/\"]+",
    ]
    start_index = text.find("ітератур")
    end_index = text.find("Обсяг")
    literature_list = []
    if start_index <= end_index:
        if start_index != -1 and end_index != -1:
            relevant_text = text[start_index:end_index]
            for pattern in patterns:
                matches = re.finditer(pattern, relevant_text)
                for match in matches:
                    literature_list.append(match.group(0))
    return literature_list


async def parse_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"
        return process_extracted_text(full_text)
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return []


def parse_docx(file_path):
    try:
        doc = Document(file_path)
        literature_list = []
        parse = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "ітератур" in text:
                parse = True
            elif "Обсяг" in text:
                break
            elif parse:
                literature_list.append(text)
        return literature_list
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return []


def save_to_csv(filename, data):
    with open(filename, 'w', newline='', encoding='utf-8') as file:
        writer = csv.DictWriter(file, fieldnames=data[0].keys())
        writer.writeheader()
        for row in data:
            writer.writerow(row)


async def create_csv_by_department(faculty):
    filename = f'departments_{CSV_DEPARTMENTS}.csv'
    headers = [
        "Department Name", "Head", "Phone", "Email", "Link",
        "Course Title", "Course Link", "Lecturers",
        "Syllabus Links", "Moscowian Lit. Entries",
        "Ukrainian Lit. Entries", "Total Lit. Entries"
    ]

    with open(filename, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(headers)

        for dept in faculty.departments:
            # Initial entry for the department with placeholders for course details
            writer.writerow([
                dept.name, dept.head, dept.phone, dept.email, dept.link,
                "", "", "", "", "", "", ""
            ])

            for course in dept.course_list:
                lecturers = ", ".join(f"{lect.name} ({lect.link})" for lect in course.lecturers)
                syllabus_links = ", ".join(course.syllabus_links)
                moscowian_lit = course.literature_statistics.moscowian_lit_count if course.literature_statistics else 0
                ukrainian_lit = course.literature_statistics.ukrainian_lit_count if course.literature_statistics else 0
                other_lit = course.literature_statistics.other_lit_count if course.literature_statistics else 0
                total_lit = moscowian_lit + ukrainian_lit + other_lit

                writer.writerow([
                    "", "", "", "", "",  # Align under department details
                    course.title, course.link, lecturers,
                    syllabus_links, moscowian_lit, ukrainian_lit, total_lit
                ])

            # Optionally add a blank row before the next department
            writer.writerow([""] * len(headers))


async def create_csv_by_course(faculty):
    filename = f'courses_{CSV_COURSES}.csv'
    headers = [
        "Course Title", "Department Name", "Number of Moscowian Lit. Entries",
        "Number of Ukrainian Lit. Entries", "Number of All Lit. Entries", "Lecturers"
    ]

    with open(filename, mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(headers)

        for dept in faculty.departments:
            for course in dept.course_list:
                lecturers = ", ".join(lect.name for lect in course.lecturers)
                moscowian_lit = course.literature_statistics.moscowian_lit_count if course.literature_statistics else 0
                ukrainian_lit = course.literature_statistics.ukrainian_lit_count if course.literature_statistics else 0
                other_lit = course.literature_statistics.other_lit_count if course.literature_statistics else 0
                total_lit = moscowian_lit + ukrainian_lit + other_lit

                writer.writerow([
                    course.title, dept.name, moscowian_lit,
                    ukrainian_lit, total_lit, lecturers
                ])


async def fetch_course_details(course_url: str):
    """Fetch details of a specific course asynchronously."""
    course_html = await fetch_html(course_url)
    soup = BeautifulSoup(course_html, 'html.parser')

    # Example of extracting lecturers and syllabus links:
    lecturers = []
    for lecturer_tag in soup.select(
            'section.lectures table tr td.lecturer a, section.practical table tr td.teachers a'):
        name = lecturer_tag.text.strip()
        link = lecturer_tag.get('href', '')
        lecturers.append(Teacher(name=name, link=link))

    syllabus_links = [tag['href'] for tag in soup.select('section.attachments a') if tag.has_attr('href')]

    return lecturers, syllabus_links


async def fetch_courses(link):
    print(f"Fetching courses from {link}")  # Debug: Show fetch operation
    html_content = await fetch_html(link)
    soup = BeautifulSoup(html_content, 'html.parser')
    courses = []

    for a_tag in soup.select('div.item a'):
        course_url = a_tag['href']
        print(f"Fetching course details and syllabus for {course_url}")  # Debug
        lecturers, syllabus_links = await fetch_course_details(course_url)
        print(f"Syllabus links found: {syllabus_links}")  # Debug: Syllabus links fetched

        if not syllabus_links:
            print(f"No syllabus links found for {course_url}, skipping literature parsing.")  # Debug
            continue

        literature_list = await fetch_literature_list(syllabus_links)
        print(f"Literature entries parsed: {len(literature_list)}")  # Debug: Check parsed entries count

        # Aggregate literature statistics based on detected languages
        lang_stats = LangStatistics()
        for entry in literature_list:
            lang = detect_language(entry)
            if lang == 'uk':
                lang_stats.ukrainian_lit_count += 1
            elif lang == 'ru':
                lang_stats.moscowian_lit_count += 1
            else:
                lang_stats.other_lit_count += 1

        course = Course(
            title=a_tag.text.strip(),
            link=course_url,
            lecturers=lecturers,
            syllabus_links=syllabus_links,
            literature_statistics=lang_stats
        )
        courses.append(course)
        print(f"Course added: {course.title}")  # Debug: Confirm course addition
    return courses


def parse_docx(file_path):
    """
    Parses a .docx file to extract literature entries based on specific markers in the text.
    """
    try:
        doc = Document(file_path)
        literature_list = []
        parse = False
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "ітератур" in text:
                parse = True  # Start recording entries after this marker
            elif "Обсяг курсу" in text:
                break  # Stop recording entries at this point
            elif parse:
                if text:  # Add non-empty text entries
                    literature_list.append(text)
        return literature_list
    except Exception as e:
        print(f"Failed to parse DOCX file: {e}")
        return []


async def fetch_literature_list(syllabus_links):
    literature_list = []
    async with aiohttp.ClientSession() as session:
        for link in syllabus_links:
            if link.endswith('.pdf'):
                # Download and parse PDF
                file_path = await download_file(session, link)
                literature_entries = await parse_pdf(file_path)
                literature_list.extend(literature_entries)
                os.remove(file_path)
            else:
                file_path = await download_file(session, link)
                literature_entries = parse_docx(file_path)
                literature_list.extend(literature_entries)
                os.remove(file_path)
    return literature_list


async def download_file(session, url):
    retries = 3  # Define the maximum number of retries
    timeout = aiohttp.ClientTimeout(total=60)  # Define a timeout for the operation

    attempt = 0
    while attempt < retries:
        attempt += 1
        try:
            async with session.get(url, timeout=timeout) as response:
                if response.status == 200:
                    file_path = url.split('/')[-1]
                    with open(file_path, 'wb') as f:
                        # Iterate over chunks of data and write them to a file
                        async for chunk in response.content.iter_chunked(1024):
                            f.write(chunk)
                    return file_path
                else:
                    print(f"Failed to download {url}: HTTP status {response.status}")
        except aiohttp.ClientError as e:
            print(f"Attempt {attempt} - Failed to download {url}: {str(e)}")
            if attempt == retries:
                print("Max retries reached, raising exception.")
                raise
            print(f"Waiting {2 ** attempt} seconds before retrying...")
            await asyncio.sleep(2 ** attempt)  # Exponential backoff strategy


async def main():
    start_time = time.time()
    url = FACULTY_LINK
    faculty = None
    department_data = []
    # Try to fetch and process data
    try:
        html_content = await fetch_html(url)
        faculty = await process_department_response(html_content)
        if faculty:
            for department in faculty.departments:
                for course in department.course_list:
                    department_data.append({
                        'Department Name': department.name,
                        'Head': department.head,
                        'Phone': department.phone,
                        'Email': department.email,
                        'Link': department.link,
                        'Course Title': course.title,
                        'Course Link': course.link,
                        'Lecturers': ', '.join([lecturer.name for lecturer in course.lecturers]),
                        'Syllabus Links': ', '.join(course.syllabus_links),
                        'Moscowian Lit. Entries': course.literature_statistics.moscowian_lit_count if course.literature_statistics else 0,
                        'Ukrainian Lit. Entries': course.literature_statistics.ukrainian_lit_count if course.literature_statistics else 0,
                        'Total Lit. Entries': course.literature_statistics.other_lit_count if course.literature_statistics else 0
                    })
    except Exception as e:
        print(f"An error occurred during data fetching or processing: {str(e)}")

    # Save data to JSON if any departments were processed
    if faculty:
        try:
            with open('output_mmf.json', 'w', encoding='utf-8') as f:
                json.dump(asdict(faculty), f, ensure_ascii=False, indent=4)
        except Exception as e:
            print(f"An error occurred during JSON saving: {str(e)}")

    # Save data to CSV regardless of processing outcome, if there's any data
    if department_data:
        try:
            await create_csv_by_department(faculty)
            await create_csv_by_course(faculty)
        except Exception as e:
            print(f"An error occurred during CSV saving: {str(e)}")
    else:
        print("No department data was processed, skipping CSV output.")

    elapsed_time = time.time() - start_time
    print(f"Program completed in {elapsed_time:.2f} seconds")


if __name__ == "__main__":
    asyncio(run(main()))
